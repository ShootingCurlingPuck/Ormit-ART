import ast

from docx.document import Document
from docx.shared import Inches, Pt
from docx.table import _Cell

from src.constants import Font, FontSize
from src.report_utils import (
    resource_path,
    restructure_date,
    safe_get_cell,
    safe_get_table,
    safe_literal_eval,
    safe_set_text,
)

DETAILS_TABLE_INDEX = 0
COGCAP_TABLE_INDEX = 1
CONCLUSION_TABLE_INDEX = 2
HUMAN_SKILLS_START_TABLE = 4
HUMAN_SKILLS_TABLE_COUNT = 5
TECH_SKILLS_START_TABLE = 9
TECH_SKILLS_TABLE_COUNT = 5
DATA_TOOLS_TABLE_START = 15
DATA_TOOLS_ITEMS_PER_TABLE = 5
INTERESTS_TABLE_INDEX = 16
LANGUAGE_SKILLS_TABLE_INDEX = 14


def add_content_cogcaptable(doc: Document, scores_str: str) -> None:
    """Adds cognitive capacity scores."""
    table = safe_get_table(doc, COGCAP_TABLE_INDEX)
    if not table:
        return

    scores = safe_literal_eval(scores_str, [])
    if not isinstance(scores, list) or len(scores) != 6:
        print("Warning: Invalid scores data. Expected a list of 6 numbers.")
        return

    for i in range(6):
        cell = safe_get_cell(table, 1, i + 1)
        if cell:
            if i == 0:
                safe_set_text(cell, scores[i])
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.bold = True
                run.underline = True
                paragraph.alignment = 1
            else:
                safe_set_text(cell, scores[i])
                paragraph = cell.paragraphs[0]
                paragraph.alignment = 1


def add_content_cogcaptable_remark(doc: Document, cogcap_output: str) -> None:
    """Adds remarks to the cognitive capacity table."""
    if not isinstance(cogcap_output, str):
        print("Warning: cogcap_output is not a string.")
        return

    table = safe_get_table(doc, COGCAP_TABLE_INDEX)
    if not table:
        return

    remark_cell = safe_get_cell(table, 2, 1)
    if not remark_cell:
        return

    safe_set_text(remark_cell, cogcap_output)


def add_content_detailstable(doc: Document, personal_details: list[str]) -> None:
    """Adds personal details to the first table."""
    table = safe_get_table(doc, DETAILS_TABLE_INDEX)
    if not table:
        return

    if not isinstance(personal_details, list):
        print("Warning: personal_details is not a list.")
        return

    if len(personal_details) == 1 and all(isinstance(ele, str) for ele in personal_details):
        personal_details = personal_details[0].split(",")

    for row_index, row in enumerate(table.rows):
        if len(row.cells) > 1:
            first_cell_text = row.cells[0].text.strip()
            second_cell_text = row.cells[1].text.strip()

            if first_cell_text == "Name candidate" and second_cell_text == ":":
                cell = safe_get_cell(table, row_index, 2)
                safe_set_text(cell, personal_details[0] if len(personal_details) > 0 else "")

            if first_cell_text == "Date of birth" and second_cell_text == ":":
                cell = safe_get_cell(table, row_index, 2)
                safe_set_text(
                    cell, restructure_date(personal_details[1]) if len(personal_details) > 1 else ""
                )

            if first_cell_text == "Position" and second_cell_text == ":":
                cell = safe_get_cell(table, row_index, 2)
                safe_set_text(cell, personal_details[2] if len(personal_details) > 2 else "")

            if first_cell_text == "Assessment date" and second_cell_text == ":":
                cell = safe_get_cell(table, row_index, 2)
                safe_set_text(
                    cell, restructure_date(personal_details[3]) if len(personal_details) > 3 else ""
                )

            if first_cell_text == "Pool" and second_cell_text == ":":
                cell = safe_get_cell(table, row_index, 2)
                safe_set_text(cell, personal_details[4] if len(personal_details) > 4 else "")


def add_icon_to_cell(cell: _Cell, score: int | None) -> None:
    """Adds an icon based on the score to a cell.

    This function has been updated to handle None values properly, which are
    now used to represent "N/A" instead of -99.
    """
    if cell is None:
        print("Warning: add_icon_to_cell called with None cell.")
        return

    safe_set_text(cell, "")

    if score is None or not isinstance(score, int):
        try:
            score = int(score) if score is not None else None
        except (ValueError, TypeError):
            print(f"Warning: Non-integer score encountered: {score}. Using N/A.")
            run = cell.paragraphs[0].add_run("N/A")
            run.font.name = Font.MONTSERRAT_REGULAR.value
            run.font.size = Pt(FontSize.SMALL.value)
            return

    if score is None:
        run = cell.paragraphs[0].add_run("N/A")
        run.font.name = Font.MONTSERRAT_REGULAR.value
        run.font.size = Pt(FontSize.SMALL.value)
        return

    run = cell.paragraphs[0].add_run()
    if score == -1:
        run.add_picture(resource_path("resources/improvement.png"), width=Inches(0.3))
    elif score == 0:
        run.add_picture(resource_path("resources/average.png"), width=Inches(0.3))
    elif score == 1:
        run.add_picture(resource_path("resources/strong.png"), width=Inches(0.3))
    else:
        print(f"Warning: Invalid score value: {score}")


def format_datatools_output(datatools_json_string: str) -> str:
    """Formats data tools output (not used in MCP, kept for consistency)."""
    try:
        return "\n".join(
            f"- {tool}: {level}" for tool, level in ast.literal_eval(datatools_json_string).items()
        )
    except (ValueError, SyntaxError):
        return "Could not parse data tools information."
