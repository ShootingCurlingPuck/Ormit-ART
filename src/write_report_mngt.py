import ast
import logging
import os
from datetime import datetime
from typing import Any

import docx
from docx.document import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import _Cell

from src.constants import LOGGER_NAME, Gender, Program

# Import common functions from report_utils
from src.report_utils import (
    replace_and_format_header_text,
    replace_piet_in_list,
    replace_text_preserving_format,
    replacePiet,
    resource_path,
    safe_get_cell,
    safe_get_table,
    safe_literal_eval,
    safe_set_text,
    split_paragraphs_at_marker_and_style,
)
from src.write_report_common import (
    add_content_cogcaptable,
    add_content_detailstable,
    add_icon_to_cell,
)

logger = logging.getLogger(LOGGER_NAME)


# --- Constants specific to MNGT report template ---
DETAILS_TABLE_INDEX = 0
COGCAP_TABLE_INDEX = 1
CONCLUSION_TABLE_INDEX = 2
FIRST_ICONS_TABLE = 4
NUM_ICONS_TABLES = 5
ITEMS_PER_ICON_TABLE = 4


def set_font_properties(cell: _Cell) -> None:
    """Sets font properties for a cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Montserrat Light"
            run.font.size = Pt(11)
            r = run._element
            run_pr = r.rPr
            if run_pr is None:
                run_pr = OxmlElement("w:rPr")
                r.append(run_pr)
            run_fonts = OxmlElement("w:rFonts")
            run_fonts.set(qn("w:ascii"), "Montserrat Light")
            run_fonts.set(qn("w:hAnsi"), "Montserrat Light")
            run_pr.append(run_fonts)


def set_font_properties2(para) -> None:
    """Sets font properties with tabs for language skills."""
    full_text = para.text
    para.clear()
    lines = full_text.splitlines()

    for line in lines:
        words = line.split()
        if words:
            for word in words[:-1]:
                run = para.add_run(word + " ")
                run.font.name = "Montserrat Light"
                run.font.size = Pt(10)
                run.bold = False
                r = run._element
                run_pr = r.rPr or OxmlElement("w:rPr")
                r.append(run_pr)
                run_fonts = OxmlElement("w:rFonts")
                run_fonts.set(qn("w:ascii"), "Montserrat Light")
                run_fonts.set(qn("w:hAnsi"), "Montserrat Light")
                run_pr.append(run_fonts)

            if words[0] == "Dutch":
                para.add_run("\t\t")
            else:
                para.add_run("\t")

            last_word = words[-1]
            last_run = para.add_run(last_word)
            last_run.font.name = "Montserrat Light"
            last_run.font.size = Pt(10)
            last_run.bold = True
            r = last_run._element
            run_pr = last_run._element.rPr or OxmlElement("w:rPr")
            r.append(run_pr)
            run_fonts = OxmlElement("w:rFonts")
            run_fonts.set(qn("w:ascii"), "Montserrat Light")
            run_fonts.set(qn("w:hAnsi"), "Montserrat Light")
            run_pr.append(run_fonts)


def update_document(
    output_dic: dict[str, Any], name: str, assessor: str, gender: Gender, program: Program
) -> str | None:
    """Updates the Word document (MNGT version)."""
    try:
        doc = docx.Document(resource_path("resources/template.docx"))  # MNGT Template
    except Exception:
        logger.exception("Failed to open template")
        return None

    # --- Prepare Replacement Dictionary ---
    replacements = {}

    # Static replacements previously done by find_and_replace_placeholder
    replacements["***"] = name.split()[0]
    replacements["ASSESSOR"] = assessor.upper()

    # Dynamic Content replacements
    dynamic_prompts = [
        "prompt2_firstimpr",
        "prompt3_personality",
        "prompt4_cogcap_remarks",
        "prompt9_interests",  # Note: prompt4_cogcap_remarks was duplicated, removed one
    ]
    for prompt_key in dynamic_prompts:
        replacement_text = output_dic.get(prompt_key, "")
        if prompt_key in ["prompt2_firstimpr", "prompt3_personality", "prompt4_cogcap_remarks"]:
            replacement_text = replacePiet(replacement_text, name, gender)  # Apply replacePiet
        # Add to dictionary using the placeholder format {key}
        replacements[f"{{{prompt_key}}}"] = replacement_text

    # Language Skill replacements
    language_replacements_str = output_dic.get("prompt5_language", "[]")
    language_levels = safe_literal_eval(language_replacements_str, [])
    if isinstance(language_levels, list):
        language_names = ["Dutch", "French", "English"]
        for index, language_name in enumerate(language_names):
            if index < len(language_levels):
                proficiency_level = language_levels[index]
                placeholder = (
                    f"{{prompt5_language_{language_name.lower()}}}"  # Placeholder per language
                )
                replacements[placeholder] = proficiency_level
            else:
                logger.warning(
                    f"No proficiency level provided for {language_name}."
                )  # Example of console warning
                placeholder = f"{{prompt5_language_{language_name.lower()}}}"
                replacements[placeholder] = "N/A"  # Or some default

    # --- Perform ALL Text Replacements ---
    replace_text_preserving_format(doc, replacements)

    # --- Handle list prompts that may contain "Piet" ---
    # (This section remains the same, operating on _original keys)
    list_prompt_keys_original = ["prompt6a_conqual_original", "prompt6b_conimprov_original"]
    for original_key in list_prompt_keys_original:
        if original_key in output_dic:
            list_str = output_dic.get(original_key, "[]")
            list_items = safe_literal_eval(list_str, [])
            if isinstance(list_items, list):
                list_items_pietless = replace_piet_in_list(list_items, name, gender)
                output_dic[original_key] = list_items_pietless
            else:
                logger.warning(f"Could not process {original_key} as a list after eval.")
                output_dic[original_key] = []
        else:
            output_dic[original_key] = []

    # --- Content in specific locations (Tables, Icons) ---
    # These functions modify specific parts and don't use the general replacement
    add_content_detailstable(doc, [name, "", program, "", ""])  # Add details to table
    replace_and_format_header_text(doc, name)  # Format header separately
    add_content_cogcaptable(doc, output_dic.get("prompt4_cogcap_scores", "[]"))
    # language_skills function call is removed as replacement is handled above

    # Profile review (icons)
    # ... (keep existing icon logic, ensuring it uses _original if needed and evals) ...
    qual_scores_str = output_dic.get(
        "prompt7_qualscore_original", output_dic.get("prompt7_qualscore", "[]")
    )
    qual_scores = safe_literal_eval(qual_scores_str, [])
    if isinstance(qual_scores, list):
        add_icons2(doc, qual_scores)
    else:
        logger.warning("Invalid qual_scores data.")

    # --- Conclusion Table ---
    # (This section remains the same, using processed _original lists)
    conclusion(doc, 0, output_dic.get("prompt6a_conqual_original", []))
    conclusion(doc, 1, output_dic.get("prompt6b_conimprov_original", []))

    # --- Save Document ---
    current_time = datetime.now()
    formatted_time = current_time.strftime("%m%d%H%M")

    # Define output directory and ensure it exists
    output_dir = "output_reports"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Save to the output directory
    updated_doc_path = os.path.join(
        output_dir, f"Assessment Report - {name} - {formatted_time}.docx"
    )
    try:
        # Apply final paragraph splitting and styling *before* saving
        split_paragraphs_at_marker_and_style(doc)
        doc.save(updated_doc_path)
        logger.info(f"Document saved: {updated_doc_path}")
    except Exception:
        logger.exception("Failed to save document")
        return None
    else:
        return updated_doc_path


def format_datatools_output(datatools_json_string: str) -> str:
    """Formats data tools output (not used in MNGT, kept for consistency)."""
    try:
        return "\n".join(
            f"- {tool}: {level}" for tool, level in ast.literal_eval(datatools_json_string).items()
        )
    except (ValueError, SyntaxError):
        return "Could not parse data tools information."


def format_interests_output(interests_json_string: str) -> str:
    """Formats interests output (not directly used in MNGT, kept for consistency)."""
    try:
        return "\n".join(f"- {interest}" for interest in ast.literal_eval(interests_json_string))
    except (ValueError, SyntaxError):
        return "Could not parse interests information."


def add_icons2(doc: Document, list_scores: list[int]) -> None:
    """Adds icons to the profile review tables (MNGT version)."""
    if not isinstance(list_scores, list):
        logger.warning("list_scores is not a list.")  # Example of console warning
        return
    table_no_start = FIRST_ICONS_TABLE
    score_index = 0
    for table_no_offset in range(NUM_ICONS_TABLES):  # Number of tables
        table_no = table_no_start + table_no_offset
        table = safe_get_table(doc, table_no)
        if not table:
            continue  # Skip to next table

        for row_no in range(1, len(table.rows)):  # Start from row 1
            if score_index < len(list_scores):  # Check if scores remain
                cell = safe_get_cell(table, row_no, 0)  # Get the first cell
                if cell:
                    add_icon_to_cell(cell, list_scores[score_index])  # Use function
                    score_index += 1
            else:
                # If we run out of scores, add N/A for remaining cells
                cell = safe_get_cell(table, row_no, 0)
                if cell:
                    run = cell.paragraphs[0].add_run("N/A")
                    run.font.name = "Montserrat Light"
                    run.font.size = Pt(9)


def conclusion(doc: Document, column: int, list_items: list[str]) -> None:
    """Adds conclusion points (already processed list) to the specified column."""
    try:
        table = doc.tables[CONCLUSION_TABLE_INDEX]
    except IndexError:
        logger.warning(f"Could not find conclusion table at index {CONCLUSION_TABLE_INDEX}")
        return

    # Expecting list_items to be a Python list already
    if not isinstance(list_items, list):
        logger.warning(f"conclusion expected a list, got {type(list_items)}")
        return

    try:
        cell = table.cell(1, column)
        # Clear existing content first
        safe_set_text(cell, "")

        # Add each item as a separate paragraph with improved bullet formatting
        for point in list_items:
            if isinstance(point, str) or point:
                # Create paragraph with manual bullet
                paragraph = cell.add_paragraph()

                # Try to use List Bullet style if available, but don't fail if it's not
                try:
                    # Check if style exists in the document
                    if "List Bullet" in doc.styles:
                        paragraph.style = "List Bullet"
                    else:
                        # Manual bullet as fallback
                        paragraph.text = "• "
                        # Style the manual bullet
                        for run in paragraph.runs:
                            run.font.name = "Montserrat"
                            run.font.size = Pt(10)
                            run.bold = True
                except Exception:
                    logger.warning("Could not apply bullet style, using manual bullet instead.")
                    # Ensure paragraph has a bullet character
                    if not paragraph.text.startswith("•"):
                        paragraph.text = "• "

                # Add the content after the bullet (or after applying style)
                content_text = str(point) if point else ""
                if paragraph.text.startswith("•"):
                    # If we have a manual bullet, add content as a separate run
                    run = paragraph.add_run(content_text)
                else:
                    # Otherwise create a new run with the content
                    run = paragraph.add_run(content_text)

                # Apply consistent font formatting
                run.font.name = "Montserrat"
                run.font.size = Pt(10)

                # Add proper XML formatting for consistent font appearance
                r = run._element
                run_pr = r.get_or_add_rPr()
                run_fonts = OxmlElement("w:rFonts")
                run_fonts.set(qn("w:ascii"), "Montserrat")
                run_fonts.set(qn("w:hAnsi"), "Montserrat")
                run_pr.append(run_fonts)

    except IndexError:
        logger.exception(f"Could not access cell (1, {column}) in conclusion table")
    except Exception:
        logger.exception(f"Error adding conclusion to column {column}")
