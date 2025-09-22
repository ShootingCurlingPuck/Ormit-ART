import ast
import logging
import os
from abc import ABC, abstractmethod
from datetime import datetime
from typing import Any, Literal

import docx
from docx.document import Document
from docx.shared import Inches, Pt
from docx.table import _Cell

from src.constants import LOGGER_NAME, Font, FontSize, Gender, Language, Program, PromptName
from src.report_utils import (
    replace_and_format_header_text,
    replace_piet_in_list,
    replace_text_preserving_format,
    replacePiet,
    resource_path,
    safe_get_cell,
    safe_get_table,
    safe_literal_eval,
    safe_set_text,
    split_paragraphs_at_marker_and_style,
)
from src.write_report_common import add_content_cogcaptable, add_content_detailstable
from src.write_report_data import (
    add_icons_data_chief,
    add_icons_data_chief_2,
    add_icons_data_tools,
    add_interests_table,
    conclusion,
    update_language_skills_table,
)

logger = logging.getLogger(LOGGER_NAME)


COGCAP_TABLE_INDEX = 1
CONCLUSION_TABLE_INDEX = 2
DATA_TOOLS_ITEMS_PER_TABLE = 5
DATA_TOOLS_TABLE_START = 15
DETAILS_TABLE_INDEX = 0
FIRST_ICONS_TABLE = 4
HUMAN_SKILLS_START_TABLE = 4
HUMAN_SKILLS_TABLE_COUNT = 5
INTERESTS_TABLE_INDEX = 16
ITEMS_PER_ICON_TABLE = 4
LANGUAGE_SKILLS_TABLE_INDEX = 14
NUM_ICONS_TABLES = 5
TECH_SKILLS_START_TABLE = 9
TECH_SKILLS_TABLE_COUNT = 5


class ReportWriter(ABC):
    scores_to_path_mapper = {
        -1: "resources/improvement.png",
        0: "resources/average.png",
        1: "resources/strong.png",
    }

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = docx.Document(file_path)

    @abstractmethod
    def _update_document(
        self,
        output_dic: dict[PromptName, Any],
        name: str,
        assessor: str,
        gender: Gender,
        program: Program,
    ) -> str | None:
        pass

    def write_report(self) -> None:
        pass

    def _add_content_cogcaptable(self, scores: list[int]) -> Document:
        """Adds cognitive capacity scores."""
        table = safe_get_table(self.doc, COGCAP_TABLE_INDEX)
        if not table:
            return self.doc

        if len(scores) != 6:
            logger.warning("Invalid scores data. Expected a list of 6 numbers.")
            return self.doc

        for i in range(6):
            cell = safe_get_cell(table, 1, i + 1)
            if cell is None:
                continue
            safe_set_text(cell, str(scores[i]))
            paragraph = cell.paragraphs[0]
            paragraph.alignment = 1
            if i == 0:
                run = paragraph.runs[0]
                run.bold = True
                run.underline = True
        return self.doc

    def _add_content_cogcaptable_remark(self, cogcap_output: str) -> Document:
        """Adds remarks to the cognitive capacity table."""
        table = safe_get_table(self.doc, COGCAP_TABLE_INDEX)
        if not table:
            return self.doc

        remark_cell = safe_get_cell(table, 2, 1)
        if not remark_cell:
            return self.doc

        safe_set_text(remark_cell, cogcap_output)
        return self.doc

    def _add_content_detailstable(self, personal_details: list[str]) -> Document:
        """Adds personal details to the first table."""
        table = safe_get_table(self.doc, DETAILS_TABLE_INDEX)
        if not table:
            return self.doc

        if len(personal_details) == 1:
            personal_details = personal_details[0].split(",")

        while len(personal_details) < 5:
            personal_details.append("")

        cell_texts = {
            "Name candidate": personal_details[0],
            "Date of birth": personal_details[1],
            "Position": personal_details[2],
            "Assessment date": personal_details[3],
            "Pool": personal_details[4],
        }

        for row_index, row in enumerate(table.rows):
            if len(row.cells) <= 1:
                continue

            first_cell_text = row.cells[0].text.strip()
            second_cell_text = row.cells[1].text.strip()
            cell = safe_get_cell(table, row_index, 2)
            personal_detail = cell_texts.get(first_cell_text)

            if second_cell_text != ":" or cell is None or personal_detail is None:
                continue

            safe_set_text(cell, personal_detail)

        return self.doc

    def _add_icon_to_cell(self, cell: _Cell, score: Literal[-1, 0, 1] | None) -> Document:
        """Adds an icon based on the score to a cell."""
        safe_set_text(cell, "")

        if score is None:
            run = cell.paragraphs[0].add_run("N/A")
            run.font.name = Font.MONTSERRAT_REGULAR.value
            run.font.size = Pt(FontSize.SMALL.value)
            return self.doc

        run = cell.paragraphs[0].add_run()
        image_path = ReportWriter.scores_to_path_mapper[score]
        run.add_picture(resource_path(image_path), width=Inches(0.3))
        return self.doc

    def _format_datatools_output(self, datatools_json_string: str) -> str:
        """Formats data tools output (not used in MNGT, kept for consistency)."""
        try:
            return "\n".join(
                f"- {tool}: {level}"
                for tool, level in ast.literal_eval(datatools_json_string).items()
            )
        except (ValueError, SyntaxError):
            return "Could not parse data tools information."


class MngtReportWriter(ReportWriter):
    pass


class DataReportWriter(ReportWriter):
    def _update_document(
        self,
        output_dic: dict[PromptName, Any],
        name: str,
        assessor: str,
        gender: Gender,
        program: Program,
    ) -> str | None:
        """Updates the Word document."""
        try:
            doc = docx.Document(resource_path("resources/Assessment_report_Data_chiefs.docx"))
        except Exception:
            logger.exception("Failed to open template")
            return None

        # --- Prepare Replacement Dictionary ---
        replacements: dict[str, str] = {}

        # Static replacements
        replacements["***"] = name.split()[0]
        replacements["ASSESSOR"] = assessor.upper()

        # Dynamic Content replacements
        dynamic_prompts = [
            PromptName.FIRST_IMPRESSION,
            PromptName.PERSONALITY,
            PromptName.COGCAP_REMARKS,
            # Interests (prompt9) are handled separately via add_interests_table
        ]
        for prompt_key in dynamic_prompts:
            replacement_text = output_dic.get(prompt_key, "")
            if prompt_key in [
                PromptName.FIRST_IMPRESSION,
                PromptName.PERSONALITY,
                PromptName.COGCAP_REMARKS,
            ]:
                replacement_text = replacePiet(replacement_text, name, gender)
            replacements[f"{{{prompt_key}}}"] = replacement_text

        # Language Skill placeholders (assuming they exist in the Data template too)
        language_replacements_str = output_dic.get(PromptName.LANGUAGE, "[]")
        language_levels = safe_literal_eval(language_replacements_str, [])
        if isinstance(language_levels, list):
            for index, language in enumerate(Language):
                if index < len(language_levels):
                    proficiency_level = language_levels[index]
                    placeholder = f"{{prompt5_language_{language.value.lower()}}}"
                    replacements[placeholder] = proficiency_level
                else:
                    logger.warning(f"No proficiency level provided for {language}.")
                    placeholder = f"{{prompt5_language_{language.value.lower()}}}"
                    replacements[placeholder] = "N/A"

        # --- Perform ALL Text Replacements ---
        replace_text_preserving_format(doc, replacements)

        # --- Handle list prompts that may contain "Piet" ---
        # Operate on the _original JSON data for these prompts
        list_prompt_keys_original = [PromptName.CONQUAL_ORIGINAL, PromptName.CONIMPROV_ORIGINAL]
        for original_key in list_prompt_keys_original:
            if original_key in output_dic:
                list_str = output_dic.get(original_key, "[]")
                # Safely evaluate the ORIGINAL JSON string
                list_items = safe_literal_eval(list_str, [])
                if isinstance(list_items, list):
                    # Replace Piet in each list item
                    list_items_pietless = replace_piet_in_list(list_items, name, gender)
                    # Store the processed list back into the _original key
                    output_dic[original_key] = list_items_pietless  # Store the list directly
                else:
                    logger.warning(f"Could not process {original_key} as a list after eval.")
                    output_dic[original_key] = []  # Ensure it's an empty list on failure
            else:
                # Ensure the key exists even if the prompt failed, to avoid errors later
                output_dic[original_key] = []

        # --- Table/Specific Location Content ---
        add_content_detailstable(doc, [name, "", program, "", ""])
        replace_and_format_header_text(doc, name)
        add_content_cogcaptable(doc, output_dic.get(PromptName.COGCAP_SCORES, "[]"))

        # --- Add language levels to language skills table (14th table) ---
        language_replacements_str = output_dic.get(PromptName.LANGUAGE, "[]")
        # Ensure backslashes are removed before parsing
        if isinstance(language_replacements_str, str):
            language_replacements_str = language_replacements_str.replace("\\", "")
        language_levels = safe_literal_eval(language_replacements_str, [])
        update_language_skills_table(doc, language_levels)

        # --- Conclusion Table ---
        # Pass the processed list from the _original key
        conclusion(doc, 0, output_dic.get(PromptName.CONQUAL_ORIGINAL, []))
        conclusion(doc, 1, output_dic.get(PromptName.CONIMPROV_ORIGINAL, []))

        # --- Interests ---
        interests_str = output_dic.get(PromptName.INTERESTS, "")
        add_interests_table(doc, interests_str)

        # Profile review (icons)
        qual_scores_str = output_dic.get(PromptName.QUALSCORE_DATA, "[]")
        qual_scores = safe_literal_eval(qual_scores_str, [])
        if isinstance(qual_scores, list) and len(qual_scores) >= 23:
            add_icons_data_chief(doc, qual_scores[:18])
            add_icons_data_chief_2(doc, qual_scores[18:23])
        else:
            logger.warning("Invalid qual_scores data.")

        # Data tools (icons)
        data_tools_str = output_dic.get(PromptName.DATATOOLS, "[]")
        data_tools_scores = safe_literal_eval(data_tools_str, [])
        if isinstance(data_tools_scores, list):
            add_icons_data_tools(doc, data_tools_scores)
        else:
            logger.warning("Invalid data_tools_scores data.")

        # --- Save Document ---
        current_time = datetime.now()
        formatted_time = current_time.strftime("%m%d%H%M")

        # Define output directory and ensure it exists
        output_dir = "output_reports"
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # Save to the output directory
        updated_doc_path = os.path.join(
            output_dir, f"Assessment Report - {name} - {formatted_time}.docx"
        )
        try:
            # Apply final paragraph splitting and styling *before* saving
            split_paragraphs_at_marker_and_style(doc)  # This handles the display format
            doc.save(updated_doc_path)
            logger.info(f"Document saved: {updated_doc_path}")
        except Exception:
            logger.exception("Failed to save document")
            return None
        else:
            return updated_doc_path


class IcpReportWriter(ReportWriter):
    pass
