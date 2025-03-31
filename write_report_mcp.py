import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import ast
import json
import re

# Import common functions from report_utils
from report_utils import (
    resource_path, _safe_get_table, _safe_get_cell, _safe_set_text, 
    _safe_add_paragraph, _safe_literal_eval,
    clean, strip_extra_quotes, clean_up, replacePiet, replace_piet_in_list,
    restructure_date, replace_and_format_header_text, open_file,
    replace_text_preserving_format, split_paragraphs_at_marker_and_style
)

# --- Constants specific to MCP report template ---
DETAILS_TABLE_INDEX = 0
COGCAP_TABLE_INDEX = 1
CONCLUSION_TABLE_INDEX = 2
FIRST_ICONS_TABLE = 4
NUM_ICONS_TABLES = 5
ITEMS_PER_ICON_TABLE = 4

def set_font_properties(cell):
    """Sets font properties for a cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Montserrat Light'
            run.font.size = Pt(11)
            r = run._element
            rPr = r.rPr
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.append(rPr)
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Montserrat Light')
            rFonts.set(qn('w:hAnsi'), 'Montserrat Light')
            rPr.append(rFonts)

def set_font_properties2(para):
    """Sets font properties with tabs for language skills."""
    full_text = para.text
    para.clear()
    lines = full_text.splitlines()

    for line in lines:
        words = line.split()
        if words:
            for word in words[:-1]:
                run = para.add_run(word + ' ')
                run.font.name = 'Montserrat Light'
                run.font.size = Pt(10)
                run.bold = False
                r = run._element
                rPr = r.rPr or OxmlElement('w:rPr')
                r.append(rPr)
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), 'Montserrat Light')
                rFonts.set(qn('w:hAnsi'), 'Montserrat Light')
                rPr.append(rFonts)

            if words[0] == 'Dutch':
                para.add_run('\t\t')
            else:
                para.add_run('\t')

            last_word = words[-1]
            last_run = para.add_run(last_word)
            last_run.font.name = 'Montserrat Light'
            last_run.font.size = Pt(10)
            last_run.bold = True
            r = last_run._element
            rPr = last_run._element.rPr or OxmlElement('w:rPr')
            r.append(rPr)
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Montserrat Light')
            rFonts.set(qn('w:hAnsi'), 'Montserrat Light')
            rPr.append(rFonts)

def update_document(output_dic, name, assessor, gender, program):
    """Updates the Word document (MCP version)."""
    try:
        doc = Document(resource_path('resources/template.docx'))  # MCP Template
    except Exception as e:
        print(f"Error: Failed to open template: {e}") # Example of console error
        return None

    # --- Prepare Replacement Dictionary ---
    replacements = {}

    # Static replacements previously done by find_and_replace_placeholder
    replacements['***'] = name.split()[0]
    replacements['ASSESSOR'] = assessor.upper()

    # Dynamic Content replacements
    dynamic_prompts = [
        'prompt2_firstimpr', 'prompt3_personality',
        'prompt4_cogcap_remarks', 'prompt9_interests' # Note: prompt4_cogcap_remarks was duplicated, removed one
    ]
    for prompt_key in dynamic_prompts:
      replacement_text = output_dic.get(prompt_key, "")
      if prompt_key in ['prompt2_firstimpr', 'prompt3_personality', 'prompt4_cogcap_remarks']:
          replacement_text = replacePiet(replacement_text, name, gender) # Apply replacePiet
      # Add to dictionary using the placeholder format {key}
      replacements[f"{{{prompt_key}}}"] = replacement_text

    # Language Skill replacements
    language_replacements_str = output_dic.get('prompt5_language', "[]")
    language_levels = _safe_literal_eval(language_replacements_str, [])
    if isinstance(language_levels, list):
        language_names = ["Dutch", "French", "English"]
        for index, language_name in enumerate(language_names):
            if index < len(language_levels):
                proficiency_level = language_levels[index]
                placeholder = f"{{prompt5_language_{language_name.lower()}}}" # Placeholder per language
                replacements[placeholder] = proficiency_level
            else:
                print(f"Warning: No proficiency level provided for {language_name}.") # Example of console warning
                placeholder = f"{{prompt5_language_{language_name.lower()}}}"
                replacements[placeholder] = "N/A" # Or some default

    # --- Perform ALL Text Replacements ---
    replace_text_preserving_format(doc, replacements)

    # --- Handle list prompts that may contain "Piet" ---
    # (This section remains the same, operating on _original keys)
    list_prompt_keys_original = ['prompt6a_conqual_original', 'prompt6b_conimprov_original']
    for original_key in list_prompt_keys_original:
        if original_key in output_dic:
            list_str = output_dic.get(original_key, "[]")
            list_items = _safe_literal_eval(list_str, [])
            if isinstance(list_items, list):
                list_items_pietless = replace_piet_in_list(list_items, name, gender)
                output_dic[original_key] = list_items_pietless
            else:
                print(f"Warning: Could not process {original_key} as a list after eval.")
                output_dic[original_key] = []
        else:
             output_dic[original_key] = []

    # --- Content in specific locations (Tables, Icons) ---
    # These functions modify specific parts and don't use the general replacement
    add_content_detailstable(doc, [name, "", program, "", ""]) # Add details to table
    replace_and_format_header_text(doc, name) # Format header separately
    add_content_cogcaptable(doc, output_dic.get('prompt4_cogcap_scores', "[]"))
    # language_skills function call is removed as replacement is handled above

    # Profile review (icons)
    # ... (keep existing icon logic, ensuring it uses _original if needed and evals) ...
    qual_scores_str = output_dic.get('prompt7_qualscore_original', output_dic.get('prompt7_qualscore', "[]"))
    qual_scores = _safe_literal_eval(qual_scores_str, [])
    if isinstance(qual_scores, list):
        add_icons2(doc, qual_scores)
    else:
        print(f"Warning: Invalid qual_scores data.")

    # --- Conclusion Table ---
    # (This section remains the same, using processed _original lists)
    conclusion(doc, 0, output_dic.get('prompt6a_conqual_original', []))
    conclusion(doc, 1, output_dic.get('prompt6b_conimprov_original', []))

    # --- Save Document ---
    current_time = datetime.now()
    formatted_time = current_time.strftime("%m%d%H%M")
    updated_doc_path = f"Assessment Report - {name} - {formatted_time}.docx"
    try:
        # Apply final paragraph splitting and styling *before* saving
        split_paragraphs_at_marker_and_style(doc)
        doc.save(updated_doc_path)
        print(f"Document saved: {updated_doc_path}") # Added print statement
        return updated_doc_path
    except Exception as e:
        print(f"Error: Failed to save document: {e}") # Example of console error
        return None

def format_datatools_output(datatools_json_string):
    """Formats data tools output (not used in MCP, kept for consistency)."""
    try:
        return "\n".join(f"- {tool}: {level}" for tool, level in ast.literal_eval(datatools_json_string).items())
    except (ValueError, SyntaxError):
        return "Could not parse data tools information."

def format_interests_output(interests_json_string):
    """Formats interests output (not directly used in MCP, kept for consistency)."""
    try:
        return "\n".join(f"- {interest}" for interest in ast.literal_eval(interests_json_string))
    except (ValueError, SyntaxError):
        return "Could not parse interests information."

def add_content_detailstable(doc, personal_details):
    table = _safe_get_table(doc, DETAILS_TABLE_INDEX)
    if not table:
        return

    if not isinstance(personal_details, list):
        print(f"Warning: personal_details is not a list.") # Example of console warning
        return

    if len(personal_details) == 1 and all(isinstance(ele, str) for ele in personal_details):
        personal_details = personal_details[0].split(',')

    for row_index, row in enumerate(table.rows):
        if len(row.cells) > 1:
            first_cell_text = row.cells[0].text.strip()
            second_cell_text = row.cells[1].text.strip()

            if first_cell_text == "Name candidate" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)
                _safe_set_text(cell, personal_details[0] if len(personal_details) > 0 else '')

            if first_cell_text == "Date of birth" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)
                _safe_set_text(cell, restructure_date(personal_details[1]) if len(personal_details) > 1 else '')

            if first_cell_text == "Position" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)
                _safe_set_text(cell, personal_details[2] if len(personal_details) > 2 else '')

            if first_cell_text == "Assessment date" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)
                _safe_set_text(cell, restructure_date(personal_details[3]) if len(personal_details) > 3 else '')

            if first_cell_text == "Pool" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)
                _safe_set_text(cell, personal_details[4] if len(personal_details) > 4 else '')

def add_content_cogcaptable(doc, scores_str):
    """Adds cognitive capacity scores."""
    table = _safe_get_table(doc, COGCAP_TABLE_INDEX)
    if not table:
        return

    scores = _safe_literal_eval(scores_str, [])
    if not isinstance(scores, list) or len(scores) != 6:
        print(f"Warning: Invalid scores data. Expected a list of 6 numbers.") # Example of console warning
        return

    for i in range(6):
        cell = _safe_get_cell(table, 1, i + 1)  # Row 1 (second row)
        if cell:
            if i == 0:
                _safe_set_text(cell, scores[i])
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.bold = True
                run.underline = True
                paragraph.alignment = 1
            else:
                _safe_set_text(cell, scores[i])
                paragraph = cell.paragraphs[0]
                paragraph.alignment = 1

def add_content_cogcaptable_remark(doc, cogcap_output):
    """Adds remarks to the cognitive capacity table."""
    if not isinstance(cogcap_output, str):
        print(f"Warning: cogcap_output is not a string.") # Example of console warning
        return

    table = _safe_get_table(doc, COGCAP_TABLE_INDEX)
    if not table:
        return

    remark_cell = _safe_get_cell(table, 2, 1)  # Row 2 (third row)
    if not remark_cell:
        return

    _safe_set_text(remark_cell, cogcap_output)

def add_icons2(doc, list_scores):
    """Adds icons to the profile review tables (MCP version)."""
    if not isinstance(list_scores, list):
        print(f"Warning: list_scores is not a list.") # Example of console warning
        return
    table_no_start = FIRST_ICONS_TABLE
    score_index = 0
    for table_no_offset in range(NUM_ICONS_TABLES): # Number of tables
      table_no = table_no_start + table_no_offset
      table = _safe_get_table(doc, table_no)
      if not table:
          continue  # Skip to next table

      for row_no in range(1, len(table.rows)): #Start from row 1
        if score_index < len(list_scores): # Check if scores remain
            cell = _safe_get_cell(table, row_no, 0) # Get the first cell
            if cell:
                add_icon_to_cell(cell, list_scores[score_index]) # Use function
                score_index += 1
        else:
            # If we run out of scores, add N/A for remaining cells
            cell = _safe_get_cell(table, row_no, 0)
            if cell:
                run = cell.paragraphs[0].add_run("N/A")
                run.font.name = 'Montserrat Light'
                run.font.size = Pt(9)

def add_icon_to_cell(cell, score):
    """
    Adds an icon based on the score to a cell (modified for MCP).
    
    This function has been updated to handle None values properly, which are
    now used to represent "N/A" instead of -99.
    """
    if cell is None:
        print(f"Warning: add_icon_to_cell called with None cell.") # Example of console warning
        return
    
    _safe_set_text(cell, "")
    
    # Handle None (N/A) or non-integer scores
    if score is None or not isinstance(score, int):
        try:
            # Try to convert to int if possible (but not if None)
            score = int(score) if score is not None else None
        except (ValueError, TypeError):
            # If it's N/A or cannot be converted, use a default icon or text
            print(f"Warning: Non-integer score encountered: {score}. Using N/A.")
            run = cell.paragraphs[0].add_run("N/A")
            run.font.name = 'Montserrat Light'
            run.font.size = Pt(9)
            return

    # Special handling for None (our placeholder for N/A)
    if score is None:
        run = cell.paragraphs[0].add_run("N/A")
        run.font.name = 'Montserrat Light'
        run.font.size = Pt(9)
        return
        
    run = cell.paragraphs[0].add_run()
    if score == -1:
        run.add_picture(resource_path("resources/improvement.png"), width=Inches(.3))
    elif score == 0:
        run.add_picture(resource_path("resources/average.png"), width=Inches(.3))
    elif score == 1:
        run.add_picture(resource_path("resources/strong.png"), width=Inches(.3))
    else:
        print(f"Warning: Invalid score value: {score}") # Example of console warning

def conclusion(doc, column, list_items):
    """Adds conclusion points (already processed list) to the specified column."""
    try:
        table = doc.tables[CONCLUSION_TABLE_INDEX]
    except IndexError:
        print(f"Warning: Could not find conclusion table at index {CONCLUSION_TABLE_INDEX}")
        return

    # Expecting list_items to be a Python list already
    if not isinstance(list_items, list):
        print(f"Warning: conclusion expected a list, got {type(list_items)}")
        return

    try:
        cell = table.cell(1, column)
        # Clear existing content first
        _safe_set_text(cell, "")

        # Add each item as a separate paragraph with bullet formatting
        for point in list_items:
            if isinstance(point, str):
                # Add pseudo-bullet for visual consistency within the table cell
                _safe_add_paragraph(cell, f'•  {point}')
            elif point:
                _safe_add_paragraph(cell, f'•  {str(point)}')

    except IndexError:
        print(f"Warning: Could not access cell (1, {column}) in conclusion table")
    except Exception as e:
        print(f"Error adding conclusion to column {column}: {e}")