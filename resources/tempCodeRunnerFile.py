import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import ast
import json
import re

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except AttributeError:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

### Initial clean:
def clean(text):
    if not text or isinstance(text, list):  # If empty or already a list
        return text
    elif isinstance(text, str):  # If string
        # Remove any square bracket items and return their contents (cleaning some classic GPT style aspects)
        cleaned_text = re.sub(r'\【.*?\】', '', text).strip()
        cleaned_text = cleaned_text.replace("```python", "")
        cleaned_text = cleaned_text.replace("```", "")
        cleaned_text = cleaned_text.replace("**", "")
        cleaned_text = strip_extra_quotes(cleaned_text)
        return cleaned_text  # Return cleaned string immediately if no lists to extract

def strip_extra_quotes(input_string):
    # Check if the string starts and ends with double quotes
    if input_string.startswith('"') and input_string.endswith('"'):
        # Remove the extra quotes
        return input_string[1:-1]
    return input_string

#######################  Actual Work  ##########################
def clean_up(loc_dic):
    with open(loc_dic, 'r') as json_file:
        loaded_data = json.load(json_file)
    results_list = {}
    for key in loaded_data.keys():
        results_list[key] = clean(loaded_data[key]) #Extract list from string
    return results_list

### Put it in Word:
#Icons
template = resource_path('resources/Assessment_report_Data_chiefs.docx')
image_path_improv = resource_path("resources/improvement.png")
image_path_average = resource_path("resources/average.png")
image_path_strong = resource_path("resources/strong.png")

def update_document(output_dic, name, assessor, gender, program):
    doc = Document(template)
    ### CONTENT
    #Personal details
    add_content_detailstable(doc, [name, "", program, "", ""]) # Program added here
    replace_and_format_header_text(doc, name)
    replace_placeholder_in_docx(doc, '***', name.split()[0], font_name='Montserrat Light')

    #Add assessor
    replace_placeholder_in_docx(doc, 'ASSESSOR', assessor, font_name='Montserrat Light')

    # First impression - always included
    if "prompt2_firstimpr" in output_dic and output_dic['prompt2_firstimpr'] != "":
        firstimpr_pietless = replacePiet(output_dic['prompt2_firstimpr'], name, gender)
        add_content_below_heading(doc, "First impression", firstimpr_pietless, "First impression")

    # Personality - always included
    if "prompt3_personality" in output_dic and output_dic['prompt3_personality'] != "":
        personality_pietless = replacePiet(output_dic['prompt3_personality'], name, gender)
        add_content_below_heading(doc, "Personality", personality_pietless, "Personality")

    # Cognitive Capacity Test results - always included
    if "prompt4_cogcap_scores" in output_dic and output_dic['prompt4_cogcap_scores'] != "":
        add_content_cogcaptable(doc, output_dic['prompt4_cogcap_scores'])

    if "prompt4_cogcap_remarks" in output_dic and output_dic['prompt4_cogcap_remarks'] != "":
        add_content_cogcaptable_remark(doc, output_dic['prompt4_cogcap_remarks'])

    # Language Skills - always included
    if "prompt5_language" in output_dic and output_dic['prompt5_language'] != "":
        language_skills(doc, output_dic['prompt5_language'])

    # Conclusion columns - always included
    if "prompt6a_conqual" in output_dic and output_dic['prompt6a_conqual'] != "":
        conclusion(doc, 0, output_dic['prompt6a_conqual'])
    if "prompt6b_conimprov" in output_dic and output_dic['prompt6b_conimprov'] != "":
        conclusion(doc, 1, output_dic['prompt6b_conimprov'])

    # Profile review - always included
    if 'prompt7_qualscore_data' in output_dic and output_dic['prompt7_qualscore_data'] != '' and 'prompt7_improvscore_data' in output_dic and output_dic['prompt7_improvscore_data'] != '':
        add_icons_data_chief(doc, output_dic['prompt7_qualscore_data'], output_dic['prompt7_improvscore_data'])

    # Data tools
    if 'prompt8_datatools' in output_dic and output_dic['prompt8_datatools'] != '' and 'prompt8_datatools' in output_dic and output_dic['prompt8_datatools'] != '':
        add_icons_data_chief_2(doc, output_dic['prompt8_datatools'], output_dic['prompt8_datatools'])

    # Interests
    if "prompt9_interests" in output_dic and output_dic['prompt9_interests'] != "":
        interests_text = format_interests_output(output_dic['prompt9_interests'])
        add_content_below_heading(doc, "Interests", interests_text, "Interests")

    # Save the document
    current_time = datetime.now()
    formatted_time = current_time.strftime("%m%d%H%M")  # Format as MMDDHHMinMin
    updated_doc_path = f"Assessment Report - {name} - {formatted_time}.docx"
    doc.save(updated_doc_path)
    os.startfile(updated_doc_path)
    return updated_doc_path

def format_datatools_output(datatools_json_string):
    try:
        datatools_dict = ast.literal_eval(datatools_json_string)
        formatted_text = ""
        for tool, level in datatools_dict.items():
            formatted_text += f"- {tool}: {level}\n"
        return formatted_text.strip() #Remove last newline
    except (ValueError, SyntaxError):
        return "Could not parse data tools information."

def format_interests_output(interests_json_string):
    try:
        interests_list = ast.literal_eval(interests_json_string)
        formatted_text = ""
        for interest in interests_list:
            formatted_text += f"- {interest}\n"
        return formatted_text.strip() #Remove last newline
    except (ValueError, SyntaxError):
        return "Could not parse interests information."


def replacePiet(text, name, gender):
    text = text.replace("Piet", name.split()[0])
    text = re.sub(r'\bthe trainee\b', name.split()[0], text, flags=re.IGNORECASE) #Replace "The trainee"

    if gender == 'M':
        replacements = {
            "She": "He",
            "she": "he",
            "Her": "Him",
            "her": "him",
            "Hers": "His",
            "hers": "his",
            "Herself": "Himself",
            "herself": "himself",
        }
        for female, male in replacements.items():
            text = re.sub(r'\b' + re.escape(female) + r'\b', male, text)

    elif gender == 'F':
        replacements = {
            "He": "She",
            "he": "she",
            "Him": "Her",
            "him": "her",
            "His": "Her",
            "his": "her",
            "Himself": "Herself",
            "himself": "herself",
        }
        for male, female in replacements.items():
            text = re.sub(r'\b' + re.escape(male) + r'\b', female, text)
    return text

def restructure_date(date_str):
    # Replace / with - to handle both delimiters
    date_str = date_str.replace('/', '-')

    try:
        # Try parsing the date in DD-MM-YYYY format
        datetime.strptime(date_str, '%d-%m-%Y')
        return date_str  # Return as is if it's already in correct format
    except ValueError:
        try:
            # Parse the input date string in YYYY-MM-DD format
            date_obj = datetime.strptime(date_str, '%Y-%m-%d')
            # Restructure and return the date in DD-MM-YYYY format
            return date_obj.strftime('%d-%m-%Y')
        except ValueError:
            return ''  # Return empty string if parsing fails

def set_font_properties(cell):
    # Set the font properties directly by modifying the cell's XML
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Montserrat Light'  # Set the font name
            run.font.size = Pt(11)  # Set font size if needed

            # Create and set font properties for Montserrat Light
            r = run._element
            rPr = r.rPr
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.append(rPr)

            rFonts = OxmlElement('w:rFonts')
            rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', 'Montserrat Light')
            rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', 'Montserrat Light')
            rPr.append(rFonts)

def set_font_properties2(para):
    # Get the full text of the paragraph
    full_text = para.text
    # Clear the existing text runs
    para.clear()

    # Split the text into lines based on new lines
    lines = full_text.splitlines()

    for line in lines:
        # Split the line into words
        words = line.split()
        # If there are words in the line
        if words:
            # Create runs for all words except the last one
            for word in words[:-1]:
                run = para.add_run(word + ' ')  # Add space after each word
                run.font.name = 'Montserrat Light'
                run.font.size = Pt(10)
                run.bold = False  # Ensure non-bold for non-last words
                r = run._element
                rPr = r.rPr

                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    r.append(rPr)

                rFonts = OxmlElement('w:rFonts')
                rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', 'Montserrat Light')
                rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', 'Montserrat Light')
                rPr.append(rFonts)

            # Add tabs before the last word based on the first word
            if words[0] == 'Dutch':
                para.add_run('\t\t')  # Three tabs for 'Dutch'
            else:
                para.add_run('\t')  # Two tabs for other cases

            last_word = words[-1]
            last_run = para.add_run(last_word)  # Last word without space
            last_run.font.name = 'Montserrat Light'
            last_run.font.size = Pt(10)
            last_run.bold = True  # Bold for last word
            r = last_run._element
            rPr = r.rPr

            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.append(rPr)

            rFonts = OxmlElement('w:rFonts')
            rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', 'Montserrat Light')
            rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', 'Montserrat Light')
            rPr.append(rFonts)

def add_content_detailstable(doc, personal_details):
    if len(personal_details)==1 and all(isinstance(ele, str) for ele in personal_details): #['allll, infoo, here']
        personal_details = personal_details[0].split(',')

    # Access the first table in the document
    table = doc.tables[0]

    # PAGE 1
    if personal_details != '':
        for row in table.rows:
            # Check the first and second cells in the current row
            if len(row.cells) > 1:  # Ensure there are at least two cells
                first_cell_text = row.cells[0].text.strip()
                second_cell_text = row.cells[1].text.strip()

                # Update for Name candidate
                if first_cell_text == "Name candidate" and second_cell_text == ":":
                    cell = row.cells[2]
                    cell.text = personal_details[0]  # Directly set the text of the cell
                    set_font_properties(cell)

                # Update for Date of Birth
                if first_cell_text == "Date of birth" and second_cell_text == ":":
                    cell = row.cells[2]
                    cell.text = restructure_date(personal_details[1])
                    set_font_properties(cell)

                # Update for Position
                if first_cell_text == "Position" and second_cell_text == ":":
                    cell = row.cells[2]
                    cell.text = personal_details[2]
                    set_font_properties(cell)

                # Update for Assessment Date
                if first_cell_text == "Assessment date" and second_cell_text == ":":
                    cell = row.cells[2]
                    cell.text = restructure_date(personal_details[3])
                    set_font_properties(cell)

                # Update for Pool
                if first_cell_text == "Pool" and second_cell_text == ":":
                    cell = row.cells[2]
                    cell.text = personal_details[4]
                    set_font_properties(cell)
    return

def add_content_below_heading(doc, heading, content, heading_name=None):
    # Split the content into paragraphs
    paragraphs = content.strip().split('\n\n')  # Split by double newline for new paragraphs

    # Find the bold heading and add content below it
    for paragraph in doc.paragraphs:
        if heading in paragraph.text:
            # Check if any run in the paragraph is bold
            for run in paragraph.runs:
                if run.bold:  # Check if the run is bold
                    # Collect the new paragraphs
                    new_paragraphs = []
                    for index, para in enumerate(paragraphs):
                        # Prepend a tab character for all except the first paragraph
                        if index > 0:
                            para = '\t' + para.strip()
                        else:
                            para = para.strip()

                        # Create a new paragraph for each segment of content
                        new_paragraphs.append(para)

                    # Insert all new paragraphs at once, in reverse order
                    for new_para in reversed(new_paragraphs):  # Reverse the order for insertion
                        # Create a new paragraph
                        new_paragraph = doc.add_paragraph(new_para)
                        # Insert the new paragraph at the correct position (after the heading)
                        doc._element.body.insert(doc._element.body.index(paragraph._element) + 1, new_paragraph._element)

                        run = new_paragraph.runs[0]
                        run.font.name = 'Montserrat Light'  # Set the font name
                        run.font.size = Pt(10)  # Set font size

                        # Create and set font properties for Montserrat Light
                        r = run._element
                        rPr = r.rPr
                        if rPr is None:
                            rPr = OxmlElement('w:rPr')
                            r.append(rPr)

                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', 'Montserrat Light')
                        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', 'Montserrat Light')
                        rPr.append(rFonts)

                    return  # Exit after adding the paragraphs
    if heading_name:
        print(f"No bold '{heading_name}' found.")


def add_content_cogcaptable(doc, scores=[]):
    if isinstance(scores, str):
        scores = scores.replace("```python", "").replace("```", "").strip()
        scores = ast.literal_eval(scores)

    if isinstance(scores, list):
        table = doc.tables[1]  # Access the second table

        for i in range(6):
            cell = table.rows[1].cells[i + 1]
            if i == 0:
                paragraph = cell.add_paragraph()
                run = paragraph.add_run(str(scores[i]))  # Convert score to string
                run.bold = True  # Set the run to bold
                run.underline = True  # Set the run to underlined
                run.font.name = 'Montserrat Light'  # Set the font name
                run.font.size = Pt(11)  # Set font size
                paragraph.alignment = 1  # Center alignment
            else:
                cell.text = str(scores[i])  # For other cells, set text normally
                set_font_properties(cell)  # Update font properties for the cell
                for paragraph in cell.paragraphs:
                    paragraph.alignment = 1  # Center alignment

def add_content_cogcaptable_remark(doc, cogcap_output):
    if isinstance(cogcap_output, str):
        cogcap_output = cogcap_output.replace("```", "").strip()  # Clean the entire text
        remark_text = cogcap_output  # Use the full cleaned text
    else:
        return  # Exit if not a string

    table = doc.tables[1]  # Access the cognitive capacity table
    remark_cell = table.rows[2].cells[1]  # Target the "Remarks" cell

    # Clear existing content
    for paragraph in remark_cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)

    # Add the full remark text
    paragraph = remark_cell.add_paragraph()
    run = paragraph.add_run(remark_text)
    run.font.name = 'Montserrat Light'
    run.font.size = Pt(10)

    return

def language_skills(doc, replacements_str):
    """
    Fills in the language skills section of a Word document based on a list of replacements.

    Args:
        doc: The Word document object (docx.Document).
        replacements_str: A string representing a list of language levels, e.g., "['C2', 'A1', 'C1']".
    """
    replacements = []
    if replacements_str and isinstance(replacements_str, str):
        try:
            # Directly parse the input string as a Python list using ast.literal_eval
            replacements = ast.literal_eval(replacements_str)
        except (SyntaxError, ValueError) as e:
            print(f"Error parsing replacements string: {e}")
            print("Please ensure the replacements string is a valid Python list string, e.g., '[\"C2\", \"A1\", \"C1\"]'")
            return # Exit if parsing fails

    if not isinstance(replacements, list):
        print("Error: Replacements is not a list after parsing.")
        return

    # Define the keywords to identify the section
    section_title = "Language Skills"
    language_names = ["Dutch", "French", "English"]

    # Flag to determine if we are in the relevant section
    in_section = False
    replacement_index = 0 # Keep track of which replacement to use

    for para in doc.paragraphs:
        if section_title in para.text:
            in_section = True
        elif in_section: # Only process paragraphs within the Language Skills section
            for lang_index, name in enumerate(language_names): # Iterate with index for clarity, though index not directly used here.
                if name in para.text and ".." in para.text:
                    if replacement_index < len(replacements):
                        replacement_value = replacements[replacement_index]
                        para.text = para.text.replace("..", replacement_value)
                        set_font_properties2(para)  # Ensure the new text is formatted
                        replacement_index += 1
                        break # Move to the next paragraph after replacing for one language
                    else:
                        print(f"Warning: No more replacements left for language: {name}")
                        break # No more replacements, move to next paragraph
            if para.text.strip() == "" and in_section and any(prev_para.text.strip() for prev_para in doc.paragraphs if doc.paragraphs.index(prev_para) < doc.paragraphs.index(para) and "Language Skills" in prev_para.text):
                # Check if the empty paragraph is after the "Language Skills" section started
                break # Exit section on empty paragraph after section has begun
        elif in_section and para.text.strip() == "": # Original exit condition, might be redundant with the one above.
            break


def add_icons_data_chief(doc, list_plus, list_min, items_per_table=4): # Adjusted function name and items_per_table is now 4 by default, but not used.
    if isinstance(list_plus, str):
        list_plus = ast.literal_eval(list_plus)
    if isinstance(list_min, str):
        list_min = ast.literal_eval(list_min)

    if isinstance(list_plus, list) and isinstance(list_min, list):
        list_res = [i + j for i, j in zip(list_plus, list_min)] #Add per item

        table_no_start = 4 #Table number where the icons for Data Chief start in the template.

        for i in range(len(list_res)):
            table_no = table_no_start + (i // 4)  # Table number increments every 4 items, starting from table 4 (table index 3)
            row_no = (i % 4) + 1 # Row number from 1 to 4 for each table

            if table_no <= len(doc.tables): # Check if table_no is within the range of available tables
                table = doc.tables[table_no-1] # Tables are 1-indexed in the document, but 0-indexed in doc.tables

                if row_no <= len(table.rows): # Check if row_no is within the range of rows in the table
                    cell = table.rows[row_no].cells[0]

                    # Clear existing paragraphs in the cell
                    for paragraph in cell.paragraphs:
                        paragraph.clear()
                        p = paragraph._element
                        p.getparent().remove(p)

                    # Directly add the image without creating a new paragraph
                    if list_res[i] == -1:
                        run = cell.add_paragraph().add_run()
                        run.add_picture(image_path_improv, width=Inches(.5))
                    elif list_res[i] == 0:
                        run = cell.add_paragraph().add_run()
                        run.add_picture(image_path_average, width=Inches(.5))
                    elif list_res[i] == 1:
                        run = cell.add_paragraph().add_run()
                        run.add_picture(image_path_strong, width=Inches(.5))
    return

def add_icons_data_chief_2(doc, list_plus, list_min, items_per_table=4): # Adjusted function name and items_per_table is now 4 by default, but not used.
    if isinstance(list_plus, str):
        list_plus = ast.literal_eval(list_plus)
    if isinstance(list_min, str):
        list_min = ast.literal_eval(list_min)

    if isinstance(list_plus, list) and isinstance(list_min, list):
        list_res = [i + j for i, j in zip(list_plus, list_min)] #Add per item

        table_no_start = 11 #Table number where the icons for Data Chief start in the template.

        for i in range(len(list_res)):
            table_no = table_no_start + (i // 4)  # Table number increments every 4 items, starting from table 4 (table index 3)
            row_no = (i % 4) + 1 # Row number from 1 to 4 for each table

            if table_no <= len(doc.tables): # Check if table_no is within the range of available tables
                table = doc.tables[table_no-1] # Tables are 1-indexed in the document, but 0-indexed in doc.tables

                if row_no <= len(table.rows): # Check if row_no is within the range of rows in the table
                    cell = table.rows[row_no].cells[0]

                    # Clear existing paragraphs in the cell
                    for paragraph in cell.paragraphs:
                        paragraph.clear()
                        p = paragraph._element
                        p.getparent().remove(p)

                    # Directly add the image without creating a new paragraph
                    if list_res[i] == -1:
                        run = cell.add_paragraph().add_run()
                        run.add_picture(image_path_improv, width=Inches(.5))
                    elif list_res[i] == 0:
                        run = cell.add_paragraph().add_run()
                        run.add_picture(image_path_average, width=Inches(.5))
                    elif list_res[i] == 1:
                        run = cell.add_paragraph().add_run()
                        run.add_picture(image_path_strong, width=Inches(.5))
    return

def conclusion(doc, column, list_items):
    if isinstance(list_items, str):
        list_items = ast.literal_eval(list_items)

    # Access the first properties table in the document
    table = doc.tables[2] #Conclusion table is still the third table in DATA template.
    cell = table.rows[1].cells[column]

    # Clear existing content in the cell
    for paragraph in cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)

    # Add bullet points
    for point in list_items:
        paragraph = cell.add_paragraph()
        text_run = paragraph.add_run(f'\t -{point}')  # Add the text after the bullet
        text_run.font.name = 'Montserrat Light'
        text_run.font.size = Pt(10)

#Last style improvements
def replace_and_format_header_text(doc, new_text):
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            if '***' in paragraph.text:
                # Replace the old text with new text
                paragraph.text = paragraph.text.replace('***', new_text)

                # Apply font formatting to each run in the paragraph
                for run in paragraph.runs:
                    run.font.name = 'Montserrat SemiBold'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(*(0xED, 0x6B, 0x55))  # Set the color
                    run.bold = True  # SemiBold approximation
                    run.italic = False  # Ensure the text is not italicized

                    # Ensure compatibility with Montserrat SemiBold through XML handling
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:ascii'), 'Montserrat SemiBold')
                    rFonts.set(qn('w:hAnsi'), 'Montserrat SemiBold')
                    run._element.rPr.append(rFonts)

def replace_placeholder_in_docx(doc, placeholder, replacement, font_name='Montserrat', font_size=10):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Replace the placeholder text
            inline = paragraph.runs
            for i in range(len(inline)):
                if placeholder in inline[i].text:
                    # Replace text and set font properties
                    inline[i].text = inline[i].text.replace(placeholder, replacement)
                    inline[i].font.name = font_name
                    inline[i].font.size = Pt(font_size)

def open_file(file_path):
    # Open the file automatically based on the operating system
    if os.name == 'nt':  # For Windows
        os.startfile(file_path)
    elif os.name == 'posix':  # For macOS and Linux
        os.system(f'open "{file_path}"')

if __name__ == "__main__":
    input_file = '1.json'

    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found. Please make sure '{input_file}' is in the same directory as the script or provide the correct path.")
    else:
        cleaned_data = clean_up(input_file)

        # --- Placeholder values - Replace these with actual logic if needed ---
        candidate_name = "Piet de Vries"  # Example name
        assessor_name = "John Doe"        # Example assessor
        candidate_gender = "M"            # Example gender (M/F)
        program_name = "Data Traineeship" # Example program

        updated_document_path = update_document(cleaned_data, candidate_name, assessor_name, candidate_gender, program_name)
        print(f"Word document updated and saved at: {updated_document_path}")